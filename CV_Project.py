from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)


document = Document()

# Profile picture
document.add_picture('Tim.jpeg', width=Inches(2.0))

# name, phone number, and email
fname = input('What is your first name? ')
lname = input('What is your last name? ')
speak('Hello ' + fname + 'how are you today?')
speak('What is your phone number ' + fname)
phone_number = input('What is your phone number? ')
email = input('What is your email address? ')

document.add_paragraph(
    fname + ' ' + lname + '  |  ' + phone_number + '  |  ' + email)

# about me
document.add_heading('About Me')
about_me = input('Tell me about yourself. ')
document.add_paragraph(about_me)

# Work Experience
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter your current company ')
from_date = input('From Date ')
to_date = input('To Date ')

p.add_run(company + '\t ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input('Describe your experience at ' + company + ' ')
p.add_run(experience_details)

# More places of employment
while True:
    has_more_experiences = input('Do you have more places of employment? Yes or no? ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter company ')
        from_date = input('From Date ')
        to_date = input('To Date ')

        p.add_run(company + '\t ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input('Describe your experience at ' + company + ' ')
        p.add_run(experience_details)
    else:
        break

# Skills Section
document.add_heading('Skills')
skill = input('Enter skill ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_skills = input('Do you have more skills to add? Yes or no ')
    if has_more_skills.lower() == 'yes':
        skill = input('Enter skill ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

# Footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using Tim's resume generator"




document.save('cv.docx')